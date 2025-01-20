from flask import Flask, request, render_template_string, send_file
import markdown
from docx import Document
import io
from markdown.treeprocessors import Treeprocessor
from markdown.extensions import Extension
from bs4 import BeautifulSoup

app = Flask(__name__)

# HTML templates as string (for single file application)
index_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Chat to Markdown</title>
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap" rel="stylesheet">
    <style>
        body {
            font-family: 'Roboto', sans-serif;
            margin: 0;
            padding: 0;
            background-color: #f5f5f5;
        }
        header {
            background-color: #333;
            color: white;
            padding: 20px;
            text-align: center;
        }
        main {
            padding: 20px;
        }
        h1 {
            font-size: 36px;
            margin-bottom: 20px;
        }
        textarea {
            width: 100%;
            padding: 10px;
            font-size: 16px;
            border: 1px solid #ddd;
            border-radius: 5px;
            box-sizing: border-box;
        }
        button {
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 16px;
            margin-top: 10px;
        }
        button:hover {
            background-color: #45a049;
        }
        .container {
            max-width: 800px;
            margin: 0 auto;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
            padding: 20px;
        }
    </style>
</head>
<body>
    <header>
        <h1>Chat to Markdown Converter</h1>
    </header>
    <main>
        <div class="container">
            <form action="/convert" method="POST">
                <textarea name="chat_input" rows="6" placeholder="Enter your chat message here"></textarea><br><br>
                <button type="submit">Convert to Markdown</button>
            </form>
        </div>
    </main>
</body>
</html>
"""

result_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Markdown</title>
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap" rel="stylesheet">
    <style>
        body {
            font-family: 'Roboto', sans-serif;
            margin: 0;
            padding: 0;
            background-color: #f5f5f5;
        }
        header {
            background-color: #333;
            color: white;
            padding: 20px;
            text-align: center;
        }
        main {
            padding: 20px;
        }
        .container {
            max-width: 800px;
            margin: 0 auto;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
            padding: 20px;
        }
        h1 {
            font-size: 36px;
        }
        pre {
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            white-space: pre-wrap;
            word-wrap: break-word;
            font-size: 16px;
            border: 1px solid #ddd;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
        }
        th, td {
            padding: 10px;
            border: 1px solid #ddd;
            text-align: left;
        }
        th {
            background-color: #f9f9f9;
        }
        .copy-btn {
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 16px;
            margin-top: 20px;
        }
        .copy-btn:hover {
            background-color: #45a049;
        }
        .copy-btn:active {
            background-color: #388e3c;
        }
        a {
            color: #4CAF50;
            text-decoration: none;
            font-weight: bold;
        }
        a:hover {
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <header>
        <h1>Converted Markdown</h1>
    </header>
    <main>
        <div class="container">
            <h3>Markdown Output:</h3>
            <pre>{{ md_output | safe }}</pre>
            <button class="copy-btn" onclick="copyToClipboard()">Copy to Clipboard</button>
            <br><br>
            <form action="/download" method="POST">
                <input type="hidden" name="markdown_content" value="{{ md_output | escape }}"/>
                <button class="copy-btn" type="submit">Download as DOC</button>
            </form>
            <br><br>
            <a href="/">Go Back</a>
        </div>
    </main>
    <script>
        function copyToClipboard() {
            var copyText = document.querySelector("pre");
            var range = document.createRange();
            range.selectNode(copyText);
            window.getSelection().removeAllRanges();
            window.getSelection().addRange(range);
            document.execCommand("copy");
            alert("Copied to clipboard!");
        }
    </script>
</body>
</html>
"""

# Custom markdown extension for handling tables and other elements
class TableExtension(Extension):
    def extendMarkdown(self, md):
        md.treeprocessors.register(TableProcessor(md), 'table', 25)

class TableProcessor(Treeprocessor):
    def run(self, root):
        for element in root.findall('.//table'):
            # Convert tables into docx format
            self.convert_table(element)

    def convert_table(self, table_element):
        # Create a table in the DOCX document
        table = self.doc.add_table(rows=0, cols=0)
        for row in table_element.findall('.//tr'):
            row_cells = table.add_row().cells
            for i, cell in enumerate(row.findall('.//td')):
                row_cells[i].text = cell.text

# Function to convert markdown to DOCX
def convert_markdown_to_docx(md_content):
    # Create a DOCX document
    doc = Document()
    
    # Convert the markdown to HTML first
    html_content = markdown.markdown(md_content, extensions=[TableExtension()])
    
    # Add a heading to the document
    doc.add_heading('Converted Chat to Markdown', 0)
    
    # Parse HTML and insert content (tables, headers, paragraphs, etc.)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Handling paragraphs
    for p in soup.find_all('p'):
        doc.add_paragraph(p.get_text())
    
    # Handling tables
    for table in soup.find_all('table'):
        # Determine the number of columns by looking at the first row
        rows = table.find_all('tr')
        num_columns = max(len(row.find_all(['td', 'th'])) for row in rows)  # Find the maximum column count
        
        # Create a table with the calculated number of columns
        table_doc = doc.add_table(rows=0, cols=num_columns)
        
        for row in rows:
            row_cells = table_doc.add_row().cells
            cells = row.find_all(['td', 'th'])
            for i, cell in enumerate(cells):
                if i < num_columns:  # Ensure we don't go out of range
                    row_cells[i].text = cell.get_text()

    # Save the document to a BytesIO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

    # Create a DOCX document
    doc = Document()
    
    # Convert the markdown to HTML first
    html_content = markdown.markdown(md_content, extensions=[TableExtension()])
    
    # Add a heading to the document
    doc.add_heading('Converted Chat to Markdown', 0)
    
    # Parse HTML and insert content (tables, headers, paragraphs, etc.)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Handling paragraphs
    for p in soup.find_all('p'):
        doc.add_paragraph(p.get_text())
    
    # Handling tables
    for table in soup.find_all('table'):
        table_doc = doc.add_table(rows=0, cols=0)
        for row in table.find_all('tr'):
            row_cells = table_doc.add_row().cells
            for i, cell in enumerate(row.find_all(['td', 'th'])):
                row_cells[i].text = cell.get_text()

    # Save the document to a BytesIO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Route for the home page
@app.route('/')
def index():
    return render_template_string(index_html)

# Route to handle the form submission
@app.route('/convert', methods=['POST'])
def convert_to_markdown():
    chat_input = request.form['chat_input']
    
    # Convert chat input to markdown
    md_output = markdown.markdown(chat_input, extensions=['markdown.extensions.tables'])
    
    # Render the result with the safe HTML output
    return render_template_string(result_html, md_output=md_output)

# Route to handle DOCX file download
@app.route('/download', methods=['POST'])
def download_docx():
    markdown_content = request.form['markdown_content']
    
    # Convert markdown content to DOCX
    docx_io = convert_markdown_to_docx(markdown_content)
    
    # Send the DOCX file for download
    return send_file(docx_io, as_attachment=True, download_name="converted_chat.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True)
